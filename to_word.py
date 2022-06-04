
def doctable(data, tabletitle, pathfile):
    from docx import Document
    import pandas as pd
    document = Document()
    data = pd.DataFrame(data)  # My input data is in the 2D list form
    document.add_heading(tabletitle)
    table = document.add_table(rows=(data.shape[0]), cols=data.shape[1])  # First row are table headers!
    for i, column in enumerate(data) :
        for row in range(data.shape[0]) :
            table.cell(row, i).text = str(data[column][row])
    document.save(pathfile)
